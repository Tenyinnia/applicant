import re
import spacy
import pypdfium2 as pdfium
from typing import List, Dict, Optional, Any
from collections import Counter
import docx
from datetime import datetime
import io
import os
# services/document_parser_service.py
from typing import Dict, Any
from sqlalchemy.orm import Session
from datetime import datetime
from uuid import UUID
from app.utils import dbSession
from app.database.models import Document, ParsedProfile, UserParsedCV
# Load spaCy model
try:
    nlp = spacy.load("en_core_web_md")
except:
    # Fallback to small model if medium is not available
    nlp = spacy.load("en_core_web_sm")

class UniversalResumeParser:
    def __init__(self, use_advanced_parsing: bool = True):
        """
        Initialize the parser
        
        Args:
            use_advanced_parsing: If True, use spaCy for advanced NLP parsing
        """
        self.use_advanced_parsing = use_advanced_parsing
        
        # Common section headers across different professions
        self.section_patterns = {
            "contact": r"(?i)^(?:contact\s*(?:info|details)?|personal\s*(?:info|details)|contact)$",
            "summary": r"(?i)(summary|profile\s*summary|about|overview)",
            "skills": r"(?i)(skills|skill\s*set|soft\s*skill|tech\s*skill|tech\s*stack|competencies|expertise|technical\s*skills|proficiencies|)",
            "experience": r"(?i)(experience|work\s*experience|employment|professional)",
            "education": r"(?i)^(?:education|qualifications|academic\s*background|degrees)$",
            "certifications": r"(?i)^(?:certifications|licenses|certificates|professional\s*certifications)$",
            "projects": r"(?i)^(?:projects|portfolio|selected\s*projects|project\s*experience)$",
            "languages": r"(?i)^(?:languages|language\s*skills)$",
            "awards": r"(?i)^(?:awards|honors|achievements|recognition)$",
            "interests": r"(?i)^(?:interests|hobbies|personal\s*interests)$", 
        }
        
        # Industry-specific keywords for better categorization
        self.industry_keywords = {
            "technology": [
                "software", "software engineer", "software developer",
                "programming", "coding", "python", "java", "javascript", "typescript",
                "golang", "c++", "c#", "php", "ruby", "rust",
                "django", "flask", "fastapi", "spring", "node", "express",
                "react", "vue", "angular", "nextjs", "svelte",
                "backend", "frontend", "fullstack",
                "api", "rest", "graphql", "microservices",
                "database", "sql", "postgresql", "mysql", "sqlite",
                "nosql", "mongodb", "redis", "elasticsearch",
                "cloud", "aws", "azure", "gcp", "serverless",
                "devops", "ci/cd", "docker", "kubernetes", "terraform",
                "system design", "scalability", "performance",
                "ai", "artificial intelligence", "machine learning",
                "deep learning", "nlp", "llm", "prompt engineering",
                "data structures", "algorithms", "agile", "scrum"
            ],

            "healthcare": [
                "healthcare", "medical", "hospital", "clinic",
                "nursing", "registered nurse", "practical nurse",
                "patient care", "patient safety", "care plan",
                "clinical", "clinical practice", "clinical documentation",
                "physician", "doctor", "consultant", "surgeon",
                "pharmacy", "pharmacist", "medication", "dispensing",
                "diagnosis", "treatment", "therapy", "rehabilitation",
                "public health", "community health",
                "mental health", "psychology", "psychiatry",
                "laboratory", "medical lab", "pathology",
                "health records", "ehr", "emr",
                "infection control", "first aid", "triage"
            ],

            "finance": [
                "finance", "financial services",
                "accounting", "accountant", "bookkeeping",
                "audit", "auditing", "compliance",
                "investment", "investments", "portfolio",
                "banking", "retail banking", "corporate banking",
                "financial analysis", "financial modeling",
                "budgeting", "forecasting",
                "tax", "taxation", "tax planning",
                "risk", "risk management",
                "wealth management", "asset management",
                "insurance", "actuarial",
                "corporate finance", "treasury",
                "economics", "financial reporting",
                "ifrs", "gaap"
            ],

            "education": [
                "education", "teaching", "learning",
                "teacher", "lecturer", "professor", "instructor",
                "student", "classroom", "lesson plan",
                "curriculum", "syllabus", "course design",
                "assessment", "evaluation", "grading",
                "pedagogy", "instructional design",
                "special education", "inclusive education",
                "early childhood education",
                "primary education", "secondary education",
                "higher education", "university", "college",
                "training", "professional development",
                "e-learning", "online learning", "lms"
            ],

            "marketing": [
                "marketing", "digital marketing",
                "branding", "brand strategy", "brand management",
                "advertising", "campaign", "promotion",
                "content marketing", "content creation",
                "seo", "sem", "ppc",
                "social media", "social media marketing",
                "email marketing", "newsletter",
                "growth marketing", "growth hacking",
                "analytics", "marketing analytics",
                "conversion", "lead generation",
                "customer acquisition", "customer retention",
                "market research", "consumer behavior",
                "copywriting", "storytelling"
            ],

            "engineering": [
                "engineering",
                "mechanical engineering", "civil engineering", "electrical engineering",
                "chemical engineering", "industrial engineering",
                "project engineering", "project management",
                "design engineering", "product design",
                "manufacturing", "production",
                "construction", "site engineering",
                "quality control", "quality assurance",
                "maintenance", "operations",
                "cad", "autocad", "solidworks", "revit",
                "safety engineering", "risk assessment",
                "process optimization", "lean", "six sigma"
            ],

            "sales": [
                "sales", "selling",
                "business development", "growth",
                "account management", "key accounts",
                "account executive", "sales executive",
                "revenue", "pipeline", "forecasting",
                "clients", "customers", "customer success",
                "negotiation", "closing",
                "relationship management", "crm",
                "lead generation", "prospecting",
                "b2b", "b2c", "enterprise sales",
                "retail sales", "inside sales", "outside sales"
            ],

            "design": [
                "design", "creative design",
                "graphic design", "visual design",
                "ui", "ux", "ui/ux",
                "product design", "interaction design",
                "illustration", "illustrator",
                "photoshop", "indesign", "after effects",
                "figma", "sketch", "adobe xd",
                "wireframe", "mockup", "prototype",
                "branding", "brand identity",
                "typography", "layout", "color theory",
                "design system", "accessibility"
            ]
        }
        
        # Education degree patterns
        self.degree_patterns = {
            "phd": r"\b(ph\.?d\.?|doctorate|doctoral)\b",
            "masters": r"\b(m\.?s\.?|m\.?a\.?|m\.?sc\.?|m\.?eng\.?|master\'?s?)\b",
            "bachelors": r"\b(b\.?s\.?|b\.?a\.?|b\.?sc\.?|b\.?eng\.?|bachelor\'?s?)\b",
            "associate": r"\b(a\.?s\.?|a\.?a\.?|associate)\b",
            "diploma": r"\b(diploma|certificate|national\s*diploma|higher\s*national\s*diploma|ND|HND)\b"
        }
        
        # Date patterns for experience extraction
        self.date_patterns = [
            r'(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4})\s*[-–—]\s*(?P<end>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4}|Present|Current|Now)',
            r'(?P<start>\d{4})\s*[-–—]\s*(?P<end>\d{4}|Present|Current|Now)',
            r'(?i)(since|from)\s+(?P<start>\d{4}|\w+\s+\d{4})(?:\s+(?:to|until)\s+(?P<end>\d{4}|present|current|now))?'
        ]
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from different file formats using pypdfium2 for PDFs
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text as string
        """
        text = ""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                # Use pypdfium2 for PDF extraction
                pdf = pdfium.PdfDocument(file_path)
                
                # Extract text from all pages
                for i in range(len(pdf)):
                    page = pdf[i]
                    textpage = page.get_textpage()
                    page_text = textpage.get_text_range()
                    text += page_text + "\n"
                
                # Close the document
                pdf.close()
                
            elif file_ext == '.docx':
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                
            elif file_ext == '.doc':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                    
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            # Try alternative method for PDF
            if file_ext == '.pdf':
                text = self._fallback_pdf_extraction(file_path)
        
        return text
    
    def _fallback_pdf_extraction(self, file_path: str) -> str:
        """Fallback method for PDF extraction"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except:
            return ""
    
    def preprocess_text(self, text: str) -> str:
        """
        Clean and normalize text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Normalize spaces BUT preserve line breaks
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{2,}', '\n\n', text)

        
        # Fix common OCR/PDF extraction issues
        text = re.sub(r'•|\uf0b7|\uf0a7', '-', text)  # Replace bullet characters
        text = re.sub(r'\s*[|]\s*', ' | ', text)  # Normalize separators
        text = re.sub(r'\s*[-–—]\s*', ' - ', text)  # Normalize dashes
        
        # Remove excessive line breaks within sentences
        text = re.sub(r'([a-z])\.\s+([A-Z])', r'\1. \2', text)
        
        return text.strip()
    
    def extract_text_with_formatting(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text with formatting information (useful for layout analysis)
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary with text and formatting info
        """
        result = {
            "raw_text": "",
            "pages": [],
            "fonts": set(),
            "font_sizes": set()
        }
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            pdf = pdfium.PdfDocument(file_path)
            
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                textpage = page.get_textpage()
                
                # Get text with positions
                text = textpage.get_text_range()
                
                # Try to get font information
                try:
                    # Get text page objects for more detailed info
                    segments = textpage.get_text_segments()
                    page_fonts = set()
                    page_sizes = set()
                    
                    for segment in segments:
                        if hasattr(segment, 'font_name'):
                            page_fonts.add(segment.font_name)
                        if hasattr(segment, 'font_size'):
                            page_sizes.add(segment.font_size)
                    
                    result["fonts"].update(page_fonts)
                    result["font_sizes"].update(page_sizes)
                    
                except:
                    pass
                
                result["raw_text"] += text + "\n"
                result["pages"].append({
                    "page_number": page_num + 1,
                    "text": text,
                    "height": page.get_height(),
                    "width": page.get_width()
                })
            
            pdf.close()
        
        else:
            # For non-PDF files, just extract text
            result["raw_text"] = self.extract_text(file_path)
        
        return result
    
    def identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify and extract sections from resume text
        
        Args:
            text: Cleaned resume text
            
        Returns:
            Dictionary with section names as keys and content as values
        """
        sections = {section: "" for section in self.section_patterns.keys()}
        sections["other"] = []
        
        lines = text.split('\n')
        current_section = "other"
        buffer = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            section_found = False
            for section_name, pattern in self.section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    # Save previous section content
                    if current_section != "other" and buffer:
                        sections[current_section] = ' '.join(buffer).strip()
                    elif current_section == "other" and buffer:
                        sections["other"].append(' '.join(buffer).strip())
                    
                    # Start new section
                    current_section = section_name
                    buffer = []
                    section_found = True
                    break
            
            if not section_found:
                buffer.append(line)
        
        # Save the last section
        if current_section != "other" and buffer:
            sections[current_section] = ' '.join(buffer).strip()
        elif current_section == "other" and buffer:
            sections["other"].append(' '.join(buffer).strip())
        
        # Clean up sections
        for section in sections:
            if isinstance(sections[section], str):
                sections[section] = sections[section].strip()
        
        return sections
    
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """
        Extract contact information from resume text
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with contact information
        """
        contact_info = {
            "name": None,
            "email": None,
            "phones": [],
            "location": None,
            "linkedin": None,
            "github": None,
            "portfolio": None,
            "other_links": []
        }
        
        # Extract emails
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info["email"] = emails[0]
            contact_info["other_emails"] = emails[1:] if len(emails) > 1 else []
        
        # Extract phone numbers (international support)
        phone_patterns = [
            r'\+\d{1,3}[\s\-]?\(?\d{1,4}\)?[\s\-]?\d{1,4}[\s\-]?\d{1,4}[\s\-]?\d{1,4}',  # International
            r'\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}',  # US/Canada
            r'\d{2,4}[\s\-]\d{6,10}',  # Other formats
            r'\d{10,15}'  # Plain numbers
        ]
        
        all_phones = []
        for pattern in phone_patterns:
            all_phones.extend(re.findall(pattern, text))
        
        # Clean phone numbers
        cleaned_phones = []
        for phone in all_phones:
            # Remove non-digit characters except leading +
            cleaned = re.sub(r'[^\d\+]', '', phone)
            if 10 <= len(cleaned) <= 15:
                cleaned_phones.append(cleaned)
        
        contact_info["phones"] = list(set(cleaned_phones))
        
        # Extract URLs and social profiles
        url_pattern = r'(https?://[^\s]+)'
        urls = re.findall(url_pattern, text)
        
        for url in urls:
            url_lower = url.lower()
            if 'linkedin.com' in url_lower or 'linked.in' in url_lower:
                contact_info["linkedin"] = url
            elif 'github.com' in url_lower:
                contact_info["github"] = url
            elif any(domain in url_lower for domain in ['portfolio', 'behance', 'dribbble', 'personal']):
                contact_info["portfolio"] = url
            else:
                contact_info["other_links"].append(url)
        
        # Extract name using NER or pattern matching
        if self.use_advanced_parsing:
            doc = nlp(text[:500])  # Only process first 500 chars for name
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    # Check if it looks like a name (has at least 2 parts)
                    parts = ent.text.split()
                    if len(parts) >= 2 and all(len(part) > 1 for part in parts):
                        contact_info["name"] = ent.text
                        break
        
        # Fallback: look for name patterns at beginning of text
        if not contact_info["name"]:
            first_lines = text.split('\n')[:5]
            for line in first_lines:
                line = line.strip()
                # Simple heuristic: line with 2-4 words, starting with capital letters
                words = line.split()
                if 2 <= len(words) <= 4:
                    if all(word[0].isupper() for word in words if word):
                        contact_info["name"] = line
                        break
        
        # Extract location
        if self.use_advanced_parsing:
            doc = nlp(text[:1000])
            for ent in doc.ents:
                if ent.label_ in ["GPE", "LOC"]:
                    contact_info["location"] = ent.text
                    break
        
        return contact_info
    
    def extract_experience(self, text: str) -> Dict[str, Any]:
        """
        Extract work experience information
        
        Args:
            text: Experience section text or full text
            
        Returns:
            Dictionary with experience information
        """
        experiences = []
        
        # Split into lines for easier processing
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Check if this line might be an experience entry
            is_experience_entry = False
            date_range = None
            
            # Look for date patterns
            context = " ".join(lines[max(0, i-2):min(len(lines), i+3)])
            for pattern in self.date_patterns:
                match = re.search(pattern, context, re.IGNORECASE)
                if match:
                    date_range = match.groupdict()
                    is_experience_entry = True
                    break
            
            if is_experience_entry:
                # Look for title and company (check previous line)
                title = ""
                company = ""
                
                if i > 0:
                    prev_line = lines[i-1]
                    # Check if previous line might be title/company
                    if not re.search(r'\d{4}', prev_line):  # Doesn't contain year
                        # Try to split by common separators
                        parts = re.split(r'\s*[|•\-–—]\s*', prev_line)
                        if len(parts) >= 2:
                            title = parts[0].strip()
                            company = ' '.join(parts[1:]).strip()
                        else:
                            title = prev_line.strip()
                
                # Collect description (next lines until empty or next date)
                description = []
                j = i + 1
                while j < len(lines) and j < i + 10:  # Limit description lines
                    next_line = lines[j]
                    # Check if next line is a new experience entry
                    next_is_experience = False
                    for pattern in self.date_patterns:
                        if re.search(pattern, next_line, re.IGNORECASE):
                            next_is_experience = True
                            break
                    
                    if next_is_experience:
                        break
                    
                    description.append(next_line)
                    j += 1
                
                experience = {
                    "title": title,
                    "company": company,
                    "start_date": date_range.get('start', ''),
                    "end_date": date_range.get('end', ''),
                    "description": ' '.join(description),
                    "duration": self._calculate_duration(date_range.get('start'), date_range.get('end'))
                }
                
                experiences.append(experience)
                i = j  # Skip processed lines
            else:
                i += 1
        
        # Calculate total experience
        total_experience = self._calculate_total_experience(experiences)
        
        return {
            "entries": experiences,
            "total_years": total_experience,
            "count": len(experiences)
        }
    
    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract education information
        
        Args:
            text: Education section text or full text
            
        Returns:
            List of education entries
        """
        education_entries = []
        
        # Look for education entries using patterns
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Check for degree indicators
            degree_found = None
            for degree_type, pattern in self.degree_patterns.items():
                if re.search(pattern, line_lower, re.IGNORECASE):
                    degree_found = degree_type
                    break
            
            if degree_found:
                # Try to extract institution and year
                institution = ""
                year = None
                gpa = None
                
                # Look for year
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match:
                    year = year_match.group(0)
                
                # Look for GPA
                gpa_match = re.search(r'GPA\s*[:]?\s*(\d\.\d{1,2})', line, re.IGNORECASE)
                if gpa_match:
                    gpa = float(gpa_match.group(1))
                
                # Institution is usually the main part of the line
                # Remove degree and year/GPA info
                institution_line = line
                for pattern in self.degree_patterns.values():
                    institution_line = re.sub(pattern, '', institution_line, flags=re.IGNORECASE)
                institution_line = re.sub(r'\b(19|20)\d{2}\b', '', institution_line)
                institution_line = re.sub(r'GPA\s*[:]?\s*\d\.\d{1,2}', '', institution_line, flags=re.IGNORECASE)
                institution = institution_line.strip(' ,|•-–—')
                
                education_entry = {
                    "degree": line.strip(),
                    "institution": institution if institution else "Unknown",
                    "graduation_year": year,
                    "gpa": gpa,
                    "degree_level": degree_found
                }
                
                education_entries.append(education_entry)
        
        return education_entries
    
    def extract_skills(self, text: str) -> Dict[str, List[str]]:
        """
        Extract skills from text
        
        Args:
            text: Skills section text or full text
            
        Returns:
            Dictionary of categorized skills
        """
        skills = {
            "technical": [],
            "soft": [],
            "tools": [],
            "languages": [],
            "certifications": [],
            "other": []
        }
        
        # Common skill patterns
        skill_patterns = {
            "technical": [
                r'\b(python|java|c\+\+|javascript|typescript|react|angular|vue|node\.?js|express)',
                r'\b(sql|mysql|postgresql|mongodb|redis|elasticsearch|oracle)',
                r'\b(aws|azure|gcp|docker|kubernetes|terraform|ansible|jenkins)',
                r'\b(machine\s+learning|ai|data\s+science|deep\s+learning|nlp|computer\s+vision)',
                r'\b(html|css|sass|less|bootstrap|tailwind|responsive\s+design)'
            ],
            "soft": [
                r'\b(leadership|communication|teamwork|collaboration|problem\s+solving)',
                r'\b(critical\s+thinking|adaptability|creativity|innovation|time\s+management)',
                r'\b(project\s+management|agile|scrum|kanban|waterfall)',
                r'\b(presentation|public\s+speaking|negotiation|conflict\s+resolution)'
            ],
            "tools": [
                r'\b(git|github|gitlab|jira|confluence|slack|teams|zoom|notion)',
                r'\b(excel|word|powerpoint|outlook|sharepoint|office\s+365)',
                r'\b(photoshop|illustrator|figma|sketch|adobe\s+xd|autocad|solidworks)'
            ],
            "languages": [
                r'\b(english|spanish|french|german|chinese|mandarin|japanese|korean|arabic|hindi)'
            ]
        }
        
        # Extract skills using patterns
        for category, patterns in skill_patterns.items():
            found_skills = set()
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                found_skills.update([match.lower() for match in matches])
            skills[category] = list(found_skills)
        
        # Look for certification patterns
        cert_patterns = [
            r'\b(AWS\s+Certified|Azure\s+Certified|Google\s+Cloud\s+Certified)',
            r'\b(PMP|PMI|Scrum\s+Master|CSM|PSM|SAFe)',
            r'\b(CPA|CFA|FRM|Series\s+\d+)',
            r'\b(CISSP|CEH|Security\+|Network\+|A\+)\b'
        ]
        
        certs = set()
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certs.update(matches)
        skills["certifications"] = list(certs)
        
        # Extract other skills (noun phrases, capitalized terms)
        if self.use_advanced_parsing:
            doc = nlp(text)
            for chunk in doc.noun_chunks:
                if 1 <= len(chunk.text.split()) <= 3:
                    skill_text = chunk.text.lower()
                    # Check if not already in other categories
                    if not any(skill_text in skill_list for skill_list in skills.values()):
                        skills["other"].append(skill_text)
        
        # Remove duplicates within categories
        for category in skills:
            skills[category] = list(set(skills[category]))
        
        return skills
    
    def _calculate_duration(self, start: str, end: str) -> str:
        """
        Calculate duration between two dates
        
        Args:
            start: Start date string
            end: End date string
            
        Returns:
            Formatted duration string
        """
        try:
            # Parse dates
            if not start or not end:
                return "N/A"
            
            # Clean date strings
            start = start.strip()
            end = end.strip()
            
            # Convert to datetime
            start_date = self._parse_date(start)
            if not start_date:
                return "N/A"
            
            if end.lower() in ['present', 'current', 'now']:
                end_date = datetime.now()
            else:
                end_date = self._parse_date(end)
                if not end_date:
                    return "N/A"
            
            # Calculate duration in months
            months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
            
            if months <= 0:
                return "N/A"
            
            # Format duration
            years = months // 12
            remaining_months = months % 12
            
            if years > 0 and remaining_months > 0:
                return f"{years} years {remaining_months} months"
            elif years > 0:
                return f"{years} years"
            else:
                return f"{months} months"
                
        except Exception as e:
            return "N/A"
    
    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """
        Parse date string to datetime object
        
        Args:
            date_str: Date string
            
        Returns:
            datetime object or None
        """
        date_str = date_str.strip()
        
        # Try different formats
        formats = [
            '%b %Y',  # "Jan 2020"
            '%B %Y',  # "January 2020"
            '%m/%Y',  # "01/2020"
            '%Y-%m',  # "2020-01"
            '%Y',     # "2020"
        ]
        
        for fmt in formats:
            try:
                return datetime.strptime(date_str, fmt)
            except:
                continue
        
        return None
    
    def _calculate_total_experience(self, experiences: List[Dict]) -> float:
        """
        Calculate total years of experience
        
        Args:
            experiences: List of experience entries
            
        Returns:
            Total years of experience
        """
        total_months = 0
        
        for exp in experiences:
            duration = exp.get("duration", "")
            # Extract months from duration string
            months = 0
            
            # Look for years
            year_match = re.search(r'(\d+)\s*years?', duration, re.IGNORECASE)
            if year_match:
                months += int(year_match.group(1)) * 12
            
            # Look for months
            month_match = re.search(r'(\d+)\s*months?', duration, re.IGNORECASE)
            if month_match:
                months += int(month_match.group(1))
            
            total_months += months
        
        return round(total_months / 12, 1)
    
    def categorize_profession(self, text: str) -> Dict[str, float]:
        """
        Categorize profession based on text content
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with profession categories and confidence scores
        """
        scores = {}
        text_lower = text.lower()
        
        for industry, keywords in self.industry_keywords.items():
            score = 0
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    score += 1
            
            # Normalize score
            normalized_score = score / len(keywords) if keywords else 0
            scores[industry] = round(normalized_score, 2)
        
        # Sort by score
        sorted_scores = dict(sorted(scores.items(), key=lambda x: x[1], reverse=True))
        
        # Get primary category
        primary = max(scores, key=scores.get) if scores else "general"
        
        return {
            "primary": primary if scores[primary] > 0.1 else "general",
            "scores": sorted_scores,
            "confidence": scores[primary] if primary in scores else 0
        }
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Main method to parse resume from file
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Parsed resume data
        """
        # Extract text
        text = self.extract_text(file_path)
        if not text:
            raise ValueError(f"Could not extract text from {file_path}")
        
        # Preprocess text
        cleaned_text = self.preprocess_text(text)
        
        # Identify sections
        sections = self.identify_sections(cleaned_text)
        
        # Extract information
        contact_info = self.extract_contact_info(cleaned_text)
        
        # Use experience section if available, otherwise full text
        experience_text = sections.get("experience", cleaned_text)
        experience_info = self.extract_experience(experience_text)
        
        # Use education section if available
        education_text = sections.get("education", cleaned_text)
        education_info = self.extract_education(education_text)
        
        # Use skills section if available
        skills_text = sections.get("skills", cleaned_text)
        skills_info = self.extract_skills(skills_text)
        
        # Categorize profession
        profession_info = self.categorize_profession(cleaned_text)
        
        # Extract summary
        summary = sections.get("summary", "").strip()
        if not summary:
            # Use first 200 characters as summary
            summary = cleaned_text[:200].strip()
            if len(cleaned_text) > 200:
                summary += "..."
        
        # Compile result
        result = {
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_type": os.path.splitext(file_path)[1],
                "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
                "parse_timestamp": datetime.now().isoformat(),
                "text_length": len(cleaned_text),
                "language": "en"  # Could be detected with langdetect
            },
            "personal_info": contact_info,
            "profession": profession_info,
            "summary": summary,
            "experience": experience_info,
            "education": education_info,
            "skills": skills_info,
            "sections": {
                section: content[:500] + "..." if isinstance(content, str) and len(content) > 500 else content
                for section, content in sections.items()
                if content and (isinstance(content, str) and content.strip() or isinstance(content, list) and content)
            }
        }
        
        return result
    
    def prepare_for_job_matching(self, parsed_data: Dict) -> Dict[str, Any]:
        """
        Prepare parsed data for job matching
        
        Args:
            parsed_data: Parsed resume data
            
        Returns:
            Data optimized for job matching
        """
        # Flatten skills
        all_skills = []
        for category, skills in parsed_data["skills"].items():
            all_skills.extend([f"{skill} ({category})" for skill in skills])
        
        # Get highest education
        highest_education = ""
        education_levels = {"phd": 5, "masters": 4, "bachelors": 3, "associate": 2, "diploma": 1}
        
        for edu in parsed_data.get("education", []):
            level = edu.get("degree_level", "")
            if level in education_levels:
                if not highest_education or education_levels[level] > education_levels.get(highest_education, 0):
                    highest_education = level
        
        # Extract keywords from summary and experience
        keywords_text = parsed_data["summary"] + " " + " ".join(
            [exp.get("description", "") for exp in parsed_data["experience"].get("entries", [])]
        )
        
        if self.use_advanced_parsing:
            doc = nlp(keywords_text)
            keywords = [token.lemma_.lower() for token in doc 
                       if not token.is_stop and token.is_alpha and len(token.text) > 2]
            keyword_counts = Counter(keywords)
            top_keywords = [kw for kw, _ in keyword_counts.most_common(20)]
        else:
            # Simple keyword extraction
            words = re.findall(r'\b[a-zA-Z]{3,}\b', keywords_text.lower())
            word_counts = Counter(words)
            top_keywords = [word for word, _ in word_counts.most_common(20)]
        
        # Compile matching data
        matching_data = {
            "candidate_id": f"cand_{hash(str(parsed_data))}",
            "name": parsed_data["personal_info"].get("name", "Unknown"),
            "title": parsed_data["experience"].get("entries", [{}])[0].get("title", "") if parsed_data["experience"].get("entries") else "",
            "total_experience_years": parsed_data["experience"].get("total_years", 0),
            "highest_education": highest_education,
            "primary_profession": parsed_data["profession"].get("primary", "general"),
            "profession_confidence": parsed_data["profession"].get("confidence", 0),
            "skills": all_skills[:50],  # Limit to top 50 skills
            "skill_count": len(all_skills),
            "keywords": top_keywords,
            "location": parsed_data["personal_info"].get("location", ""),
            "email": parsed_data["personal_info"].get("email", ""),
            "phone": parsed_data["personal_info"].get("phones", [""])[0] if parsed_data["personal_info"].get("phones") else "",
            "summary": parsed_data["summary"][:300],  # Truncate summary
            "last_updated": parsed_data["metadata"]["parse_timestamp"],
            "experience_count": parsed_data["experience"].get("count", 0),
            "education_count": len(parsed_data.get("education", [])),
            "certifications": parsed_data["skills"].get("certifications", [])
        }
        
        return matching_data


# # Utility function for quick parsing
# def parse_resume_file(file_path: str, return_matching_data: bool = False) -> Dict:
#     """
#     Quick utility function to parse a resume file
    
#     Args:
#         file_path: Path to resume file
#         return_matching_data: If True, return data optimized for job matching
        
#     Returns:
#         Parsed resume data
#     """
#     parser = UniversalResumeParser()
    
#     try:
#         parsed_data = parser.parse_resume(file_path)
        
#         if return_matching_data:
#             return parser.prepare_for_job_matching(parsed_data)
#         else:
#             return parsed_data
            
#     except Exception as e:
#        pass

resume_parser = UniversalResumeParser()

class DocumentParsingService:
    @staticmethod
    async def parse_document_async(db: Session, document_id: str, file_path: str):
        """Async wrapper for document parsing"""
        # In production, you might want to use Celery or BackgroundTasks
        # For now, we'll do it synchronously but you can adapt
        return DocumentParsingService.parse_document(db, document_id, file_path)
    
    @staticmethod
    def parse_document(db: Session, document_id: str, file_path: str):
        """Parse a document and save results to database"""
        from models import Document
        
        try:
            # Get the document
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                return {"error": "Document not found"}
            
            # Parse the document
            parsed_data = resume_parser.parse_resume(file_path)
            
            # Update document with parsed data
            document.parsed_data = parsed_data
            document.parsed_at = datetime.now()
            document.is_parsed = True
            
            # Extract quick-access fields
            if "name" in parsed_data:
                document.candidate_name = parsed_data.get("name")
            if "email" in parsed_data:
                document.candidate_email = parsed_data.get("email")
            if "experience_years" in parsed_data:
                document.total_experience = parsed_data.get("experience_years")
            if "skills" in parsed_data:
                document.skills_summary = parsed_data.get("skills")
            
            db.commit()
            
            return {
                "success": True,
                "document_id": document_id,
                "parsed_data": parsed_data
            }
            
        except Exception as e:
            # Log error and update document
            document.parsed_data = {"error": str(e)}
            document.is_parsed = False
            document.parsing_error = str(e)
            db.commit()
            
            return {
                "success": False,
                "error": str(e),
                "document_id": document_id
            }
    
    @staticmethod
    def get_parsed_data_for_matching(db: Session, user_id: str):
        """Get parsed data optimized for job matching"""
        from app.database.models import Document
        
        # Get user's latest CV (assuming one active CV per user)
        document = db.query(Document).filter(
            Document.user_id == user_id,
            Document.file_type.in_(['cv', 'resume']),
            Document.is_parsed == True
        ).order_by(Document.created_at.desc()).first()
        
        if not document or not document.parsed_data:
            return None
        
        # Extract matching-ready data
        parsed = document.parsed_data
        
        return {
            "candidate_id": str(document.user_id),
            "document_id": str(document.id),
            "name": parsed.get("name"),
            "email": parsed.get("email"),
            "phone": parsed.get("phone"),
            "summary": parsed.get("summary"),
            "skills": parsed.get("skills", []),
            "experience_years": parsed.get("experience_years", 0),
            "keywords": parsed.get("skills", []),  # Can be enhanced
            "parsed_at": document.parsed_at.isoformat() if document.parsed_at else None
        }
    
def parse_file(file_path: str) -> Dict[str, Any]:
    """Return structured payload or raise ValueError."""
    data = resume_parser.parse_resume(file_path)
    if not data.get("is_parsed"):
        raise ValueError(data.get("error", "Unknown parse error"))
    return resume_parser.prepare_for_job_matching(data)




def parse_cv_task(user_id: UUID, document_id: UUID, file_path: str):
    db_gen = dbSession()  # ✅ FIX 1: real session instance
    db = next(db_gen)
    try:
        parser = UniversalResumeParser()

        raw_text = parser.extract_text(file_path)
        clean_text = parser.preprocess_text(raw_text)

        sections = parser.identify_sections(clean_text)
        contact = parser.extract_contact_info(clean_text)
        experience = parser.extract_experience(
            sections.get("experience", clean_text)
        )
        education = parser.extract_education(
            sections.get("education", clean_text)
        )

        payload = {
            "contact": contact,
            "summary": sections.get("summary"),
            "skills": sections.get("skills"),
            "experience": experience,
            "education": education,
            "raw_sections": sections,
        }

        # 🔄 UPSERT ParsedProfile
        parsed = (
            db.query(ParsedProfile)
            .filter(ParsedProfile.document_id == document_id)
            .first()
        )

        if parsed:
            parsed.payload = payload
        else:
            parsed = ParsedProfile(
                user_id=user_id,
                document_id=document_id,
                payload=payload,
            )
            db.add(parsed)

        # 🔗 mark this CV as user's current parsed CV
        db.merge(
            UserParsedCV(
                user_id=user_id,
                document_id=document_id,
            )
        )

        db.commit()  # ✅ works now

    except Exception as e:
        db.rollback()
        print(f"CV parsing failed: {e}")
        raise
        

    finally:
        db.close()  # ✅ safe now
        try:
            next(db_gen, None)
        except:
            pass